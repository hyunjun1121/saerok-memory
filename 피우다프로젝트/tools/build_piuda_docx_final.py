from __future__ import annotations

import datetime as dt
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_TXT = ROOT / "지원서_양식.txt"
ASSET_DIR = ROOT / "application_assets"
OUT_DOCX = ROOT / "2026_글로벌_피우다프로젝트_지원신청서_새록정원.docx"

PRODUCT_NAME = "새록정원"
OLD_PRODUCT_NAME = "기억정원"


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:cs"), "Calibri")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size_pt: float = 10.2, color: str = "000000", bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size_pt=size_pt, bold=bold, color=color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 140, bottom: int = 90, end: int = 140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge, color, size in (
        ("top", "9E9E9E", "6"),
        ("left", "9E9E9E", "6"),
        ("bottom", "9E9E9E", "6"),
        ("right", "9E9E9E", "6"),
        ("insideH", "D0D0D0", "4"),
        ("insideV", "D0D0D0", "4"),
    ):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def clear_cell(cell) -> None:
    for paragraph in list(cell.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)


def add_cell_paragraph(cell, text: str = "", style: str | None = None, bold: bool = False, color: str = "000000", size: float = 10.2):
    paragraph = cell.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.18
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size_pt=size, bold=bold, color=color)
    return paragraph


def read_sections() -> list[tuple[str, str]]:
    text = SOURCE_TXT.read_text(encoding="utf-8")
    chunks = re.split(r"^##\s+\d+\.\s+(.+?)\s*$", text, flags=re.MULTILINE)
    sections: list[tuple[str, str]] = []
    for i in range(1, len(chunks), 2):
        title = chunks[i].strip()
        body = chunks[i + 1].strip()
        body = body.replace(OLD_PRODUCT_NAME, PRODUCT_NAME).replace("`", "")
        body = normalize_body(title, body)
        if title == "개발물명":
            body = add_name_rationale(body)
        sections.append((title, body))
    return sections


def normalize_body(title: str, body: str) -> str:
    blocks = [b.strip() for b in re.split(r"\n\s*\n", body) if b.strip() and b.strip() != "---"]
    normalized: list[str] = []
    for block in blocks:
        if title == "개발물 설명" and block.startswith("현재 확인 가능한 화면 자산은"):
            normalized.append("대표 화면은 본문 하단의 화면 예시에 정리했습니다. 홈, 학습 피드백, 개인 기억 선택, 정원 보상, 가족·보호자 안내, 설정 및 삭제 흐름을 통해 현재 구현물의 실제 사용 장면을 확인할 수 있습니다.")
            continue
        if title == "기대효과" and block.startswith("출처 목록은 제출 전"):
            normalized.append("출처 목록은 아래와 같습니다. 접근일은 2026년 5월 15일입니다.")
            continue
        normalized.append(block)
    return "\n\n".join(normalized)


def add_name_rationale(body: str) -> str:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", body) if p.strip()]
    if paragraphs:
        paragraphs[0] = PRODUCT_NAME
    rationale = (
        "명칭은 ‘새록새록’ 떠오르는 기억과 학습 결과가 정원처럼 자라나는 화면 구조를 함께 담기 위해 "
        f"‘{PRODUCT_NAME}’으로 정리했습니다. 질병명을 전면에 내세우지 않아 고령 사용자가 낙인감 없이 접근할 수 있고, "
        "비의료적 취미·회상 루틴이라는 개발 방향에도 부합합니다."
    )
    return "\n\n".join(paragraphs[:2] + [rationale] + paragraphs[2:])


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.2)

    for style_name, size, color, bold in (
        ("Title", 18, "1F4D78", True),
        ("Subtitle", 11, "4A5568", False),
        ("Heading 1", 14, "2E74B5", True),
        ("Heading 2", 12, "2E74B5", True),
        ("Caption", 8.5, "555555", False),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold

    return doc


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("2026 글로벌 피우다프로젝트 지원신청서 작성본")
    set_run_font(run, size_pt=18, bold=True, color="1F4D78")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"개발물명: {PRODUCT_NAME}")
    set_run_font(run, size_pt=11, color="4A5568")

    note_table = doc.add_table(rows=1, cols=1)
    note_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(note_table)
    cell = note_table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, 120, 180, 120, 180)
    clear_cell(cell)
    add_cell_paragraph(
        cell,
        "원본 신청서의 표 항목에 맞춰 DOCX에서 편집하기 쉽도록 재구성한 작성본입니다. "
        "제출 전 팀원 실명, 권역, 날짜, 서명, NIPA 교육 증빙 여부는 실제 정보로 교체해야 합니다.",
        size=9.5,
        color="333333",
    )
    doc.add_paragraph()


def add_section_table(doc: Document, title: str, body: str) -> None:
    if title == "개발주제 관련 팀 및 팀원 역량":
        doc.add_page_break()
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(3.0)
    table.columns[1].width = Cm(14.0)
    set_table_borders(table)

    label, value = table.rows[0].cells
    label.width = Cm(3.0)
    value.width = Cm(14.0)
    for cell in (label, value):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_margins(cell)

    set_cell_shading(label, "F4F6F9")
    clear_cell(label)
    p = add_cell_paragraph(label, title, bold=True, color="1F4D78", size=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    clear_cell(value)
    for block in [b.strip() for b in re.split(r"\n\s*\n", body) if b.strip()]:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if lines and all(line.startswith("- ") for line in lines):
            for line in lines:
                p = add_cell_paragraph(value, line[2:].strip(), style="List Bullet")
                set_paragraph_font(p, size_pt=9.9)
            continue
        if lines and all(re.match(r"^\d+\.\s+", line) for line in lines):
            for line in lines:
                item = re.sub(r"^\d+\.\s+", "", line).strip()
                p = add_cell_paragraph(value, item, style="List Number")
                set_paragraph_font(p, size_pt=9.5)
            continue
        text = " ".join(lines)
        if text.startswith("※"):
            p = add_cell_paragraph(value, text, color="7A5A00", size=9.4)
        else:
            p = add_cell_paragraph(value, text, bold=(text == PRODUCT_NAME), size=10.0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()


def add_images(doc: Document) -> None:
    images = [
        ("final_01_home.png", "그림 1. 홈 화면과 학습 시작"),
        ("final_03_lesson_feedback.png", "그림 2. 선택형 학습과 정답 피드백"),
        ("final_04_memory_selection.png", "그림 3. 개인 기억 주제 선택"),
        ("final_06_garden.png", "그림 4. 정원 보상 화면"),
        ("final_07_family.png", "그림 5. 가족·보호자 안내 화면"),
        ("final_08_settings.png", "그림 6. 언어 설정 및 기억 카드 삭제"),
    ]
    available = [(ASSET_DIR / name, caption) for name, caption in images if (ASSET_DIR / name).exists()]
    if not available:
        return

    doc.add_page_break()
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("화면 예시")
    set_paragraph_font(heading, size_pt=14, bold=True, color="2E74B5")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(5)
    run = intro.add_run("아래 이미지는 현재 구현물의 대표 화면입니다. 제출용 최종본에서는 심사 분량에 맞춰 일부 이미지만 남길 수 있습니다.")
    set_run_font(run, size_pt=9.8)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, (path, caption) in enumerate(available):
        if i and i % 2 == 0:
            table.add_row()
        row = table.rows[i // 2]
        cell = row.cells[i % 2]
        set_cell_margins(cell, 100, 120, 100, 120)
        clear_cell(cell)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(7.2))
        cap = add_cell_paragraph(cell, caption, style="Caption", size=8.5, color="555555")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{PRODUCT_NAME} | 2026 글로벌 피우다프로젝트 지원신청서")
    set_run_font(run, size_pt=8, color="666666")


def build() -> None:
    doc = setup_document()
    add_footer(doc)
    add_title_block(doc)
    for title, body in read_sections():
        add_section_table(doc, title, body)
    add_images(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
