from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(r"C:\project\saerok-memory")
FINAL = ROOT / "피우다프로젝트" / "final"
SCREENSHOTS = ROOT / "피우다프로젝트" / "application_assets" / "final_qa" / "ko"
REPORT_ASSETS = FINAL / "report_assets"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.com")

DETAIL_MD = FINAL / "Haru_종합_연구근거_구현보고서.md"
DETAIL_DOCX = FINAL / "Haru_종합_연구근거_구현보고서.docx"
BROAD_MD = FINAL / "Haru_큰틀_종합보고서.md"
BROAD_DOCX = FINAL / "Haru_큰틀_종합보고서.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text.strip())
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(9)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_image_alt(paragraph, alt_text: str) -> None:
    drawings = paragraph._p.xpath(".//w:drawing")
    if not drawings:
        return
    doc_pr = drawings[-1].xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("descr", alt_text)


def crop_asset(source_name: str, output_name: str, top: int, bottom: int) -> None:
    source = SCREENSHOTS / source_name
    if not source.exists():
        return
    REPORT_ASSETS.mkdir(parents=True, exist_ok=True)
    output = REPORT_ASSETS / output_name
    with Image.open(source) as image:
        width, height = image.size
        safe_top = max(0, min(top, height - 1))
        safe_bottom = max(safe_top + 1, min(bottom, height))
        image.crop((0, safe_top, width, safe_bottom)).save(output)


def prepare_report_assets() -> None:
    crop_asset("22_report-caregiver.png", "22_report-caregiver_top.png", 0, 1900)
    crop_asset("21_report-counselor.png", "21_report-counselor_top.png", 0, 2200)


def resolve_image_path(image_name: str) -> Path:
    report_path = REPORT_ASSETS / image_name
    if report_path.exists():
        return report_path
    return SCREENSHOTS / image_name


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.name = "Malgun Gothic"
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor.from_string("555555")


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_text(cell, text)
    doc.add_paragraph()


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    return text.strip()


def add_markdown_table(doc: Document, rows: list[str]) -> None:
    parsed: list[list[str]] = []
    for row in rows:
        cells = [clean_inline(cell) for cell in row.strip().strip("|").split("|")]
        if all(set(cell.replace(":", "").replace("-", "").strip()) == set() for cell in cells):
            continue
        parsed.append(cells)

    if not parsed:
        return

    max_cols = max(len(row) for row in parsed)
    table = doc.add_table(rows=len(parsed), cols=max_cols)
    table.style = "Table Grid"
    for row_index, row in enumerate(parsed):
      for col_index in range(max_cols):
        text = row[col_index] if col_index < len(row) else ""
        cell = table.cell(row_index, col_index)
        set_cell_text(cell, text, bold=row_index == 0)
        if row_index == 0:
            set_cell_shading(cell, "F4F6F9")
    doc.add_paragraph()


def add_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    table_buffer: list[str] = []
    in_code = False

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_table()
            in_code = not in_code
            continue

        if in_code:
            if stripped:
                paragraph = doc.add_paragraph(stripped)
                paragraph.style = doc.styles["Normal"]
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(stripped)
            continue

        flush_table()

        if not stripped or stripped == "---":
            continue

        if stripped.startswith("# "):
            doc.add_heading(clean_inline(stripped[2:]), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(clean_inline(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(clean_inline(stripped[4:]), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(clean_inline(stripped[5:]), level=3)
        elif stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(clean_inline(stripped[2:]))
        elif re.match(r"^\d+\.\s+", stripped):
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.add_run(clean_inline(re.sub(r"^\d+\.\s+", "", stripped)))
        else:
            paragraph = doc.add_paragraph(clean_inline(stripped))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    flush_table()


def add_figure(doc: Document, image_name: str, caption: str, alt_text: str, width: float = 3.0) -> None:
    image_path = resolve_image_path(image_name)
    if not image_path.exists():
        return
    table = doc.add_table(rows=1, cols=1)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    set_image_alt(paragraph, alt_text)
    caption_paragraph = cell.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption_paragraph.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("555555")
    doc.add_paragraph()


def write_markdown_files() -> None:
    source_parts = [
        FINAL / "Haru_연구근거_구현신빙성_보고서.md",
        FINAL / "research_claim_matrix.md",
        FINAL / "implementation_qa_report.md",
        FINAL / "final_validation_log.md",
    ]
    detail_text = [
        "# Haru 종합 연구근거 구현보고서",
        "",
        "본 문서는 Haru의 현재 구현체를 기준으로 인지 루틴, 보호자·상담사 리포트, Haru 자체 종합 주의 신호가 어떤 연구적 배경과 제품 설계 판단에 따라 구현되었는지 정리한다.",
        "",
        "Haru는 공식 인지검사 문항이나 점수 체계를 복제하지 않는다. 대신 신뢰 가능한 문헌과 공개 자료에서 확인되는 인지 영역을 Haru식 짧은 일상 루틴으로 재설계하고, 반복 활동 기록과 보호자 관찰을 조합해 대화 준비에 필요한 참고 신호를 제공한다.",
        "",
    ]
    for path in source_parts:
        if path.exists():
            detail_text.append(f"## {path.stem}")
            detail_text.append(path.read_text(encoding="utf-8"))
            detail_text.append("")
    DETAIL_MD.write_text("\n".join(detail_text), encoding="utf-8")

    broad_text = """# Haru 큰틀 종합보고서

Haru는 고령 사용자가 하루 한 번 짧은 인지·기억 루틴을 수행하고, 그 과정에서 쌓이는 활동 기록과 개인 기억 단서를 가족 및 상담사가 이해하기 쉬운 리포트로 연결하는 서비스이다.

## 1. Haru가 해결하려는 문제

고령자의 일상 변화는 한 번의 검사나 한 장의 설문으로만 이해하기 어렵다. 기억, 주의, 언어, 날짜 감각, 그리기와 같은 활동은 매일의 생활 속에서 조금씩 달라질 수 있고, 가족이 느끼는 변화도 함께 살펴볼 때 상담 준비가 쉬워진다.

## 2. 서비스의 기본 구조

Haru는 학습 홈, 일일 레슨, 결과, 기억 정원, 보호자·상담사 리포트, 설정 화면으로 구성된다. 사용자는 부담이 적은 선택형 활동을 수행하고, 보호자와 상담사는 공유 허용된 기억과 활동 흐름을 바탕으로 다음 대화를 준비한다.

## 3. 인지 루틴의 의미

Haru의 루틴은 지연 단어 회상, 숫자 기억, 범주 말하기, 주의 전환, 색상 집중, 날짜 감각, 도형 그리기처럼 넓은 인지 영역을 짧게 경험하도록 설계되었다. 이는 공식 검사를 복제하는 방식이 아니라, 매일 반복하기 쉬운 원본 활동으로 재구성한 것이다.

## 4. 개인 기억의 역할

사용자는 자신의 이야기를 구체적인 단서로 남길 수 있다. 사람, 장소, 감정, 이야기 요약은 이후 회상 질문과 가족 대화 소재로 이어진다. 개인 기억은 기본적으로 비공개이며, 사용자가 공유를 허용한 내용만 보호자·상담사 화면에서 활용된다.

## 5. Haru 자체 종합 주의 신호

Haru는 반복 루틴 참여, 회상 기록, 반응 흐름, 그리기 활동, 보호자 관찰을 하나의 공식 점수로 합치지 않는다. 대신 여러 약한 신호를 함께 살펴 `안정적`, `살펴보기`, `대화 필요` 수준의 Haru 자체 참고 신호로 정리한다. 이 신호는 가족과 상담사가 어떤 이야기를 먼저 꺼내면 좋을지 결정하는 데 초점을 둔다.

## 6. 보호자와 상담사를 위한 리포트

보호자 화면은 가족이 느낀 일상 변화를 간단히 기록하도록 돕는다. 상담사 화면은 최근 활동, 공유 기억, 보호자 관찰, Haru advisory 신호를 한 화면에 모아 상담 전 대화 주제를 준비하게 한다.

## 7. 개인정보와 해석 경계

Haru는 로컬 저장소 기반 MVP로 작동하며, 기억 카드는 기본적으로 비공개이다. Haru의 참고 신호는 공식 진단이나 공식 검사 점수가 아니라, 반복 활동과 관찰 기록을 바탕으로 한 대화 준비 자료이다.

## 8. 현재 완성도

현재 구현체는 한국어·일본어·영어 화면을 지원하고, 26개 테스트 파일의 79개 테스트를 통과했다. Playwright로 3개 언어 각 23장, 총 69장의 화면 캡처도 재검증했다.
"""
    BROAD_MD.write_text(broad_text, encoding="utf-8")


def build_detail_docx() -> None:
    doc = Document()
    configure_document(doc)
    add_title(
        doc,
        "Haru 종합 연구근거 구현보고서",
        "인지 루틴, 보호자·상담사 리포트, Haru 자체 종합 주의 신호 구현 근거",
    )
    add_callout(
        doc,
        "요약: Haru는 공식 검사를 복제하지 않고, 연구 문헌에서 확인되는 넓은 인지 영역을 일상 루틴으로 재설계한다. 현재 구현은 반복 활동 기록과 보호자 관찰을 결합해 상담 전 대화를 준비하는 Haru 자체 참고 신호를 제공한다.",
    )
    for image_name, caption, alt, width in [
        ("01_home.png", "그림 1. 홈 화면의 일일 루틴 및 Haru advisory 안내", "Haru 홈 화면과 advisory 안내 카드", 3.0),
        ("16_lesson-delayed-word-recall.png", "그림 2. 지연 단어 회상 확인 화면", "단어 회상 선택형 확인 화면", 2.4),
        ("14_lesson-shape-copy.png", "그림 3. 도형 복사 그리기 화면", "손가락으로 도형을 따라 그리는 화면", 2.4),
        ("21_report-counselor_top.png", "그림 4. 상담사 리포트와 Haru 종합 주의 신호", "상담사 리포트 상단의 Haru 종합 주의 신호", 1.45),
    ]:
        add_figure(doc, image_name, caption, alt, width)
    add_markdown(doc, DETAIL_MD.read_text(encoding="utf-8"))
    doc.save(DETAIL_DOCX)


def build_broad_docx() -> None:
    doc = Document()
    configure_document(doc)
    add_title(
        doc,
        "Haru 큰틀 종합보고서",
        "누구나 이해할 수 있는 서비스 의미, 기능 구조, 연구 기반 요약",
    )
    add_callout(
        doc,
        "Haru는 매일의 짧은 활동을 통해 사용자가 기억과 감정을 되돌아보고, 가족과 상담사가 더 나은 대화를 준비하도록 돕는 일상 인지·회상 서비스이다.",
    )
    add_markdown(doc, BROAD_MD.read_text(encoding="utf-8"))
    for image_name, caption, alt, width in [
        ("01_home.png", "그림 1. 하루 한 번 시작하는 루틴 화면", "Haru 홈 화면", 3.0),
        ("17_lesson-memory-story.png", "그림 2. 개인 기억 이야기를 남기는 화면", "개인 기억 이야기 입력 화면", 2.1),
        ("20_garden.png", "그림 3. 활동이 보상으로 이어지는 기억 정원", "Haru 기억 정원 화면", 2.1),
        ("22_report-caregiver_top.png", "그림 4. 보호자 관찰 메모와 요약 신호", "보호자 리포트 상단의 관찰 메모와 요약 신호", 1.55),
        ("21_report-counselor_top.png", "그림 5. 상담사 대화 준비 리포트", "상담사 리포트 상단의 대화 준비 정보", 1.45),
        ("23_settings.png", "그림 6. 언어와 로컬 데이터 관리", "설정 화면", 2.1),
    ]:
        add_figure(doc, image_name, caption, alt, width)
    doc.save(BROAD_DOCX)


def convert_to_pdf(docx_path: Path) -> Path | None:
    if not SOFFICE.exists():
        return None
    subprocess.run(
        [
            str(SOFFICE),
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    pdf_path = docx_path.with_suffix(".pdf")
    return pdf_path if pdf_path.exists() else None


def main() -> None:
    prepare_report_assets()
    write_markdown_files()
    build_detail_docx()
    build_broad_docx()
    detail_pdf = convert_to_pdf(DETAIL_DOCX)
    broad_pdf = convert_to_pdf(BROAD_DOCX)
    print(f"DETAIL_DOCX={DETAIL_DOCX}")
    print(f"DETAIL_PDF={detail_pdf}")
    print(f"BROAD_DOCX={BROAD_DOCX}")
    print(f"BROAD_PDF={broad_pdf}")


if __name__ == "__main__":
    main()
